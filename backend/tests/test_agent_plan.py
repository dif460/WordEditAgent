import pytest
from docx import Document

from agent.nodes import _build_plan
from engine.classify import classify_paragraphs
from engine.document_model import DocumentModel
from engine.reader import read_document


@pytest.fixture()
def model(tmp_path):
    path = str(tmp_path / "s.docx")
    doc = Document()
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这是一段较长的正文内容，用于测试标题识别与规划逻辑是否正确工作。")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("另一段正文内容，同样比较长，用来避免被误判成标题。")
    doc.save(path)
    return read_document(path)


def test_classify_headings(model: DocumentModel):
    cls = classify_paragraphs(model)
    levels = {pid: lvl for pid, (kind, lvl) in cls.items() if kind == "heading"}
    assert levels == {"p_0001": 1, "p_0003": 2}


def test_build_plan(model: DocumentModel):
    rules = {
        "headings": [
            {"level": 1, "font": "黑体", "size": 16, "bold": True, "alignment": "center"},
            {"level": 2, "font": "黑体", "size": 14, "bold": True},
        ],
        "body": {"font": "宋体", "size": 12, "line_spacing": 1.5, "first_line_indent": "2字符", "alignment": "justify"},
        "page": {"size": "A4", "margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17}},
    }
    plan = _build_plan(model, rules)
    tools = [s["tool"] for s in plan]
    assert "set_section_format" in tools
    assert tools.count("set_heading_style") == 2
    assert "set_run_font" in tools
    assert "set_paragraph_format" in tools
