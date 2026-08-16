"""研究论文默认格式测试：分类/规划/执行/校验闭环（无网络、无 Word COM）。"""
from __future__ import annotations

from docx import Document

from agent.nodes import _build_plan
from engine.classify import classify_paragraphs
from engine.controller import DocumentController
from rules.loader import load_rule_template, normalize_rules
from verify.rule_check import check_rules


def _make_paper(path: str) -> None:
    doc = Document()
    doc.add_paragraph("基于深度学习的图像识别研究")          # 题目
    doc.add_paragraph("计算机专业 张三")                     # 副信息
    doc.add_paragraph("摘要：本文研究深度学习在图像识别中的应用。")
    doc.add_paragraph("关键词：深度学习；图像识别；卷积神经网络")
    doc.add_paragraph("1  前言")                            # 一级标题
    doc.add_paragraph("近年来深度学习发展迅速，本节介绍研究背景。")  # 正文
    doc.add_paragraph("2.1 研究方法")                       # 二级标题
    doc.add_paragraph("2.1.1 数据集")                       # 三级标题
    doc.add_paragraph("本文使用 ImageNet 数据集进行训练。")   # 正文
    doc.add_paragraph("表1 实验结果")                       # 表标题
    doc.add_paragraph("图1 模型结构")                       # 图标题
    doc.add_paragraph("参考文献")                           # 参考文献标题
    doc.add_paragraph("[1] 邹成伟.城市停车需求预测[J].公路,2013,8(8):147-150.")
    doc.add_paragraph("[2] R. A. Rich. A plan for parking[J]. HFM, 1997,10(4):26-32.")
    doc.add_paragraph("致谢")                               # 致谢标题
    doc.add_paragraph("感谢导师的悉心指导。")
    doc.add_paragraph("附录")                               # 附录标题
    doc.add_paragraph("附录数据详见补充材料。")
    doc.save(path)


def test_classify_paper(tmp_path):
    path = str(tmp_path / "paper.docx")
    _make_paper(path)
    from engine.reader import read_document

    model = read_document(path)
    cls = classify_paragraphs(model)
    kinds = {pid: kind for pid, (kind, _) in cls.items()}
    assert kinds["p_0001"] == "title"
    assert kinds["p_0005"] == "heading" and cls["p_0005"][1] == 1
    assert kinds["p_0007"] == "heading" and cls["p_0007"][1] == 2
    assert kinds["p_0008"] == "heading" and cls["p_0008"][1] == 3
    assert kinds["p_0010"] == "table_caption"
    assert kinds["p_0011"] == "figure_caption"
    assert kinds["p_0012"] == "references_heading"
    assert kinds["p_0013"] == "references_entry"
    assert kinds["p_0015"] == "acknowledgement_heading"
    assert kinds["p_0017"] == "appendix_heading"


def test_default_format_apply_and_verify(tmp_path):
    path = str(tmp_path / "paper.docx")
    _make_paper(path)
    controller = DocumentController(path)
    model = controller.model()

    rules = normalize_rules(load_rule_template("default"))
    plan = _build_plan(model, rules)
    assert len(plan) > 0

    for step in plan:
        controller.dispatch(step["tool"], step["args"])

    out = str(tmp_path / "out.docx")
    controller.save_document(out)

    # 重新读取结果校验
    from engine.reader import read_document

    result_model = read_document(out)
    para_by_text = {p.text.strip(): p for p in result_model.paragraphs if p.text.strip()}

    # 题目：黑体 16pt
    title = para_by_text["基于深度学习的图像识别研究"]
    assert title.runs[0].east_asia == "黑体"
    assert title.runs[0].size == 16

    # 一级标题：黑体 14pt
    h1 = para_by_text["1  前言"]
    assert h1.runs[0].east_asia == "黑体"
    assert h1.runs[0].size == 14

    # 二级标题：黑体 12pt
    h2 = para_by_text["2.1 研究方法"]
    assert h2.runs[0].east_asia == "黑体"
    assert h2.runs[0].size == 12

    # 三级标题：宋体 12pt
    h3 = para_by_text["2.1.1 数据集"]
    assert h3.runs[0].east_asia == "宋体"
    assert h3.runs[0].size == 12

    # 正文：宋体 12pt，西文 Times New Roman
    body = para_by_text["本文使用 ImageNet 数据集进行训练。"]
    assert body.runs[0].east_asia == "宋体"
    assert body.runs[0].font == "Times New Roman"
    assert body.runs[0].size == 12

    # 参考文献标题：黑体 14pt
    rh = para_by_text["参考文献"]
    assert rh.runs[0].east_asia == "黑体"
    assert rh.runs[0].size == 14

    # 参考文献条目：宋体 12pt
    re1 = para_by_text["[1] 邹成伟.城市停车需求预测[J].公路,2013,8(8):147-150."]
    assert re1.runs[0].east_asia == "宋体"
    assert re1.runs[0].size == 12

    # 图表标题：黑体 10.5pt
    fig = para_by_text["图1 模型结构"]
    assert fig.runs[0].east_asia == "黑体"
    assert abs(fig.runs[0].size - 10.5) < 0.01

    # 规则校验应通过
    report = check_rules(result_model, rules)
    assert report["ok"], f"校验未通过：{report['issues']}"
