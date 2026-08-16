import os

import pytest
from docx import Document

from engine.controller import DocumentController
from engine.reader import read_document


@pytest.fixture()
def sample_docx(tmp_path):
    path = str(tmp_path / "sample.docx")
    doc = Document()
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("这是一段正文内容，用于测试格式化的效果。")
    doc.add_heading("1.1 研究背景", level=2)
    doc.add_paragraph("另一段正文。")
    doc.save(path)
    return path


def test_read_document(sample_docx):
    model = read_document(sample_docx)
    assert len(model.paragraphs) >= 4
    headings = [p for p in model.paragraphs if p.outline_level]
    assert headings[0].outline_level == 1
    assert headings[0].text == "第一章 绪论"


def test_controller_formatting(sample_docx, tmp_path):
    ctrl = DocumentController(sample_docx)
    model = ctrl.model()

    # 找到第一个标题和第一段正文
    heading = next(p for p in model.paragraphs if p.outline_level == 1)
    body = next(p for p in model.paragraphs if p.outline_level is None and p.text.strip())

    r1 = ctrl.set_heading_style(heading.id, level=1, font="黑体", east_asia="黑体", size=16, bold=True, alignment="center")
    assert r1["ok"] is True
    assert heading.id in r1["changed"]

    r2 = ctrl.set_run_font(body.id, font="宋体", east_asia="宋体", size=12)
    assert r2["ok"] is True

    r3 = ctrl.set_paragraph_format(body.id, line_spacing=1.5, first_line_indent="2字符", alignment="justify")
    assert r3["ok"] is True

    out = str(tmp_path / "out.docx")
    ctrl.save_document(out)

    # 重新读取验证
    model2 = read_document(out)
    heading2 = next(p for p in model2.paragraphs if p.outline_level == 1)
    body2 = next(p for p in model2.paragraphs if p.outline_level is None and p.text.strip())

    assert heading2.runs[0].east_asia == "黑体"
    assert heading2.runs[0].size == 16
    assert heading2.format.alignment == "center"

    assert body2.runs[0].east_asia == "宋体"
    assert body2.runs[0].size == 12
    assert body2.format.line_spacing == 1.5
    assert body2.format.first_line_indent == "2字符"
    assert body2.format.alignment == "justify"


def test_undo_last(sample_docx):
    ctrl = DocumentController(sample_docx)
    model = ctrl.model()
    body = next(p for p in model.paragraphs if p.outline_level is None and p.text.strip())
    ctrl.dispatch("set_run_font", {"paragraph_id": body.id, "size": 20})
    assert ctrl.model().paragraphs[int(body.id.split("_")[-1]) - 1].runs[0].size == 20
    res = ctrl.undo_last()
    assert res["ok"] is True
    assert ctrl.model().paragraphs[int(body.id.split("_")[-1]) - 1].runs[0].size != 20
