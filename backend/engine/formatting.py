"""基于 python-docx 的格式化操作（直接修改 docx 对象）。

所有函数均为底层原语，返回 None；由 controller 负责调用并记录变更/支持撤销。
"""
from __future__ import annotations

from typing import Optional

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "A5": (14.8, 21.0),
    "Letter": (21.59, 27.94),
    "B5": (17.6, 25.0),
}


def set_run_font(
    run,
    font: Optional[str] = None,
    east_asia: Optional[str] = None,
    size: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    f = run.font
    if font is not None:
        f.name = font
    if east_asia is not None:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        f.size = Pt(float(size))
    if bold is not None:
        f.bold = bool(bold)
    if italic is not None:
        f.italic = bool(italic)
    if color is not None:
        from docx.shared import RGBColor

        try:
            f.color.rgb = RGBColor.from_string(color.lstrip("#"))
        except Exception:
            pass


def set_paragraph_alignment(paragraph, alignment: Optional[str]) -> None:
    if alignment and alignment in ALIGN_MAP:
        paragraph.alignment = ALIGN_MAP[alignment]


def set_line_spacing(paragraph, line_spacing: Optional[float]) -> None:
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = float(line_spacing)


def clear_paragraph_indent(paragraph) -> None:
    """强制清除段落所有缩进（左缩进、首行缩进、悬挂缩进）。
    
    标题必须顶格，此函数直接移除 w:ind 元素，不受样式继承影响。
    """
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        pPr.remove(ind)


def set_space(paragraph, before: Optional[float] = None, after: Optional[float] = None) -> None:
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(float(before))
    if after is not None:
        pf.space_after = Pt(float(after))


def set_first_line_indent(paragraph, value: Optional[str], pt_size: float = 12.0) -> None:
    """设置首行缩进。value 形如 "2字符" / "0.74cm" / "0字符"。

    "0" / "0字符" 时直接移除 w:ind 元素，确保样式继承不残留缩进
    （图片、公式、图例表例等需要完全顶格居中的元素均依赖此行为）。
    """
    if not value:
        return

    import re

    v = str(value).strip()

    # "0" / "0字符" / "0字" → 彻底清除缩进，与 clear_paragraph_indent 一致
    chars_m = re.search(r"([\d.]+)\s*(?:字符|字)", v)
    if chars_m:
        chars = float(chars_m.group(1))
        if chars == 0.0:
            clear_paragraph_indent(paragraph)
            return
    elif v == "0":
        clear_paragraph_indent(paragraph)
        return

    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)

    cm_m = re.search(r"([\d.]+)\s*cm", v, re.IGNORECASE)
    if chars_m:
        chars = float(chars_m.group(1))
        ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))
        ind.set(qn("w:firstLine"), str(int(round(chars * pt_size * 20))))
    elif cm_m:
        paragraph.paragraph_format.first_line_indent = Cm(float(cm_m.group(1)))
    else:
        # 纯数值按字符处理
        try:
            chars = float(v)
            ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))
            ind.set(qn("w:firstLine"), str(int(round(chars * pt_size * 20))))
        except ValueError:
            return


def set_hanging_indent(paragraph, value: Optional[str], pt_size: float = 12.0) -> None:
    """设置悬挂缩进（参考文献条目用）。

    value 形如 "4字符"。悬挂缩进 = 首行回退 N 字符，其余行缩进 N 字符。
    在 OOXML 中用 w:hangingChars（字符单位）+ w:hanging（twips）表示。
    """
    if not value:
        return

    import re

    v = str(value).strip()
    chars_m = re.search(r"([\d.]+)\s*(?:字符|字)", v)
    if chars_m:
        chars = float(chars_m.group(1))
    else:
        try:
            chars = float(v)
        except ValueError:
            return

    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    # 清除首行缩进，避免冲突
    for attr in ("firstLine", "firstLineChars"):
        if ind.get(qn(f"w:{attr}")) is not None:
            del ind.attrib[qn(f"w:{attr}")]
    ind.set(qn("w:hangingChars"), str(int(round(chars * 100))))
    ind.set(qn("w:hanging"), str(int(round(chars * pt_size * 20))))


def apply_style(paragraph, style_name: str) -> None:
    paragraph.style = style_name


def set_outline_level(paragraph, level: Optional[int]) -> None:
    if level is None:
        return
    ppr = paragraph._p.get_or_add_pPr()
    ol = ppr.get_or_add_outlineLvl()
    ol.set(qn("w:val"), str(int(level)))


def modify_style_definition(
    doc,
    style_name: str,
    font: Optional[str] = None,
    east_asia: Optional[str] = None,
    size: Optional[float] = None,
    bold: Optional[bool] = None,
) -> bool:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return False

    if font is not None:
        style.font.name = font
    if size is not None:
        style.font.size = Pt(float(size))
    if bold is not None:
        style.font.bold = bool(bold)
    if east_asia is not None:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), east_asia)
    return True


def set_section_format(section, page_size: Optional[str] = None, margins: Optional[dict] = None) -> None:
    if page_size:
        size = PAGE_SIZES_CM.get(page_size)
        if size:
            section.page_width = Cm(size[0])
            section.page_height = Cm(size[1])
            section.orientation = WD_ORIENT.PORTRAIT
    if margins:
        if margins.get("top") is not None:
            section.top_margin = Cm(float(margins["top"]))
        if margins.get("bottom") is not None:
            section.bottom_margin = Cm(float(margins["bottom"]))
        if margins.get("left") is not None:
            section.left_margin = Cm(float(margins["left"]))
        if margins.get("right") is not None:
            section.right_margin = Cm(float(margins["right"]))


def set_header_footer(section, header: Optional[str] = None, footer: Optional[str] = None) -> None:
    if header is not None:
        p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        p.text = header
    if footer is not None:
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        p.text = footer


def add_page_number_footer(section, alignment: str = "center", font: str = "Times New Roman", size: float = 9.0) -> None:
    """在页面底端添加居中对齐的页码字段（小五号 Times New Roman）。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = ALIGN_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)

    # 清除已有内容
    for r in p.runs:
        r._element.getparent().remove(r._element)

    # 添加 PAGE 字段
    run = p.add_run()
    run.font.name = font
    run.font.size = Pt(size)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    run._r.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


# ==================== 三线表转换 ====================

def convert_to_three_line_table(table) -> None:
    """将普通表格转为标准三线表：
    - 删除左右竖线
    - 上下边框 1 磅粗实线 (w:sz="8")
    - 中间（表头下）0.75 磅细实线 (w:sz="6")
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 移除已有边框设置
    old_borders = tblPr.find(qn("w:tblBorders"))
    if old_borders is not None:
        tblPr.remove(old_borders)

    borders = OxmlElement("w:tblBorders")

    def _add_border(name: str, val: str, sz: str, color: str = "000000") -> None:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)

    _add_border("top", "single", "8")       # 1 磅
    _add_border("bottom", "single", "8")    # 1 磅
    _add_border("left", "none", "0")
    _add_border("right", "none", "0")
    _add_border("insideH", "single", "6")   # 0.75 磅
    _add_border("insideV", "none", "0")

    tblPr.append(borders)


# ==================== 标点半角清洗 ====================

_FULL_TO_HALF = {
    "\u201c": "\u201c",  # "（不变）
    "\uff0c": ",",  # ，
    "\uff1b": ";",  # ；
    "\uff1a": ":",  # ：
    "\uff08": "(",  # （
    "\uff09": ")",  # ）
    "\u201c": "\"",  # "
    "\u201d": "\"",  # "
    "\u2018": "'",   # '
    "\u2019": "'",   # '
    "\u3002": ".",   # 。
    "\uff01": "!",   # ！
    "\uff1f": "?",   # ？
}


def normalize_punctuation(paragraph) -> int:
    """将段落中中文全角标点替换为英文半角（保留中文句号），返回替换数量。"""
    import re

    count = 0
    # 注意：不转换中文句号 "。"，中文正文中句号应为全角
    full_map = {
        "\uff0c": ",", "\uff1b": ";", "\uff1a": ":", "\uff08": "(", "\uff09": ")",
        "\u201c": "\"", "\u201d": "\"", "\u2018": "'", "\u2019": "'",
        "\uff01": "!", "\uff1f": "?",
    }
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = run.text
        for full, half in full_map.items():
            if full in new_text:
                new_text = new_text.replace(full, half)
                count += 1
        if new_text != run.text:
            run.text = new_text
    return count


# ==================== 批量西文字体替换 ====================

def batch_replace_latin_font(paragraph, latin_font: str = "Times New Roman") -> int:
    """将段落中所有 ASCII 数字和英文字符的 run 字体设为 latin_font。
    返回修改的 run 数量。
    """
    import re

    count = 0
    for run in paragraph.runs:
        if not run.text:
            continue
        if re.search(r"[a-zA-Z0-9]", run.text):
            run.font.name = latin_font
            count += 1
    return count


# ==================== 分页符插入 ====================

def insert_page_break(paragraph) -> None:
    """在段落前插入分页符（w:br w:type="page"）。"""
    run = OxmlElement("w:r")
    br_elem = OxmlElement("w:br")
    br_elem.set(qn("w:type"), "page")
    run.append(br_elem)
    paragraph._p.insert(0, run)


def add_blank_paragraph_after(paragraph, doc) -> object:
    """在指定段后插入一个空段落，返回新段落。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, doc)


# ==================== 日期文本清洗 ====================

def normalize_date_text(paragraph) -> int:
    """清洗日期文本中多余空格，如 '20   25 年  6 月' -> '2025年6月'。
    返回替换数量。
    """
    import re

    count = 0
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = re.sub(r"(\d)\s+(\d)", r"\1\2", run.text)
        new_text = re.sub(r"(\d)\s+(年|月|日)", r"\1\2", new_text)
        new_text = re.sub(r"(年|月)\s+(\d)", r"\1\2", new_text)
        if new_text != run.text:
            run.text = new_text
            count += 1
    return count


# ==================== 关键词分隔符统一 ====================

def normalize_keyword_separators(paragraph) -> int:
    """将关键词段落的分隔符统一为英文半角分号 '; '。
    返回替换数量。
    """
    import re

    count = 0
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = run.text.replace("\uff1b", ";").replace("\uff0c", ";")
        new_text = re.sub(r";\s*;", ";", new_text)
        new_text = re.sub(r";(?!\s)", "; ", new_text)
        if new_text != run.text:
            run.text = new_text
            count += 1
    return count


# ==================== 图片尺寸标准化 ====================

def normalize_image_size(paragraph, max_width_cm: float = 14.0,
                          max_height_cm: float = 6.0) -> int:
    """将段落中内嵌图片尺寸标准化，不超过最大宽高。
    返回调整的图片数量。
    """
    count = 0
    for drawing in paragraph._element.findall(".//" + qn("w:drawing")):
        extent = drawing.find(".//" + qn("wp:extent"))
        if extent is not None:
            try:
                cx = int(extent.get("cx", "0"))
                cy = int(extent.get("cy", "0"))
                w_cm = cx / 360000.0
                h_cm = cy / 360000.0
                if w_cm > max_width_cm:
                    ratio = max_width_cm / w_cm
                    extent.set("cx", str(int(cx * ratio)))
                    extent.set("cy", str(int(cy * ratio)))
                    count += 1
                elif h_cm > max_height_cm:
                    ratio = max_height_cm / h_cm
                    extent.set("cx", str(int(cx * ratio)))
                    extent.set("cy", str(int(cy * ratio)))
                    count += 1
            except (ValueError, TypeError):
                pass
    return count


# ==================== 图表编号重排 ====================

def renumber_figures_tables(paragraphs: list) -> dict:
    """遍历全文，将图/表编号重排为连续递增序号。
    返回 {old_label: new_label} 映射。
    """
    import re

    fig_pattern = re.compile(r"\u56fe\s*(\d+)")  # 图
    tab_pattern = re.compile(r"\u8868\s*(\d+)")  # 表

    fig_counter = 0
    tab_counter = 0
    mapping = {}

    for p in paragraphs:
        text = p.text
        fm = fig_pattern.search(text)
        if fm:
            fig_counter += 1
            old = f"\u56fe{fm.group(1)}"
            new = f"\u56fe{fig_counter}"
            mapping[old] = new
            for run in p.runs:
                if old in (run.text or ""):
                    run.text = run.text.replace(old, new)

        tm = tab_pattern.search(text)
        if tm:
            tab_counter += 1
            old = f"\u8868{tm.group(1)}"
            new = f"\u8868{tab_counter}"
            mapping[old] = new
            for run in p.runs:
                if old in (run.text or ""):
                    run.text = run.text.replace(old, new)

    return mapping


# ==================== 公式引用上浮修正 ====================

def fix_formula_citation_superscript(paragraph) -> int:
    """检查段落中公式引用的上浮标注 (w:vertAlign)，移除非上标引用。
    返回修正数量。
    """
    count = 0
    for run in paragraph.runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr is not None:
            vert_align = rPr.find(qn("w:vertAlign"))
            if vert_align is not None and vert_align.get(qn("w:val")) == "superscript":
                rPr.remove(vert_align)
                count += 1
    return count


# ==================== 修复日志 ====================

def create_fix_log() -> dict:
    """创建修复日志容器。"""
    return {"fixes": [], "warnings": [], "errors": []}


def add_fix(log: dict, category: str, detail: str) -> None:
    log["fixes"].append({"category": category, "detail": detail})


def add_warning(log: dict, category: str, detail: str) -> None:
    log["warnings"].append({"category": category, "detail": detail})


def add_error(log: dict, category: str, detail: str) -> None:
    log["errors"].append({"category": category, "detail": detail})
