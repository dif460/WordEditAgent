"""docx -> DocumentModel 解析。"""
from __future__ import annotations

from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from engine.document_model import (
    DocumentModel,
    ParagraphFormatModel,
    ParagraphModel,
    RunModel,
    SectionModel,
    StyleModel,
    TableModel,
)

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

_LINE_RULE_MAP = {
    WD_LINE_SPACING.EXACTLY: "exact",
    WD_LINE_SPACING.AT_LEAST: "at_least",
    WD_LINE_SPACING.MULTIPLE: "multiple",
    WD_LINE_SPACING.SINGLE: "single",
}


def read_document(file_path: str) -> DocumentModel:
    doc = Document(file_path)
    return build_model(doc)


def build_model(doc: Document) -> DocumentModel:
    model = DocumentModel()

    model.sections = [_read_section(s) for s in doc.sections]
    model.styles = _read_styles(doc)

    for i, p in enumerate(doc.paragraphs):
        model.paragraphs.append(_read_paragraph(p, i))

    for i, t in enumerate(doc.tables):
        rows = len(t.rows)
        cols = len(t.columns)
        text = "\n".join(c.text for row in t.rows for c in row.cells)
        model.tables.append(TableModel(id=f"t_{i + 1:04d}", rows=rows, cols=cols, text=text[:2000]))

    return model


def _read_section(section) -> SectionModel:
    margins = {
        "top": round(section.top_margin.cm, 2) if section.top_margin else 0,
        "bottom": round(section.bottom_margin.cm, 2) if section.bottom_margin else 0,
        "left": round(section.left_margin.cm, 2) if section.left_margin else 0,
        "right": round(section.right_margin.cm, 2) if section.right_margin else 0,
    }
    page_size = "A4"
    try:
        w = section.page_width.cm
        h = section.page_height.cm
        if abs(w - 21.0) < 0.5 and abs(h - 29.7) < 0.5:
            page_size = "A4"
        elif abs(w - 21.59) < 0.5 and abs(h - 27.94) < 0.5:
            page_size = "Letter"
        else:
            page_size = f"{w:.1f}x{h:.1f}cm"
    except Exception:
        page_size = "A4"

    header = _header_footer_text(section.header)
    footer = _header_footer_text(section.footer)
    return SectionModel(margins=margins, page_size=page_size, header=header, footer=footer)


def _header_footer_text(header_footer) -> str:
    try:
        return "\n".join(p.text for p in header_footer.paragraphs).strip()
    except Exception:
        return ""


def _read_styles(doc) -> dict[str, StyleModel]:
    styles: dict[str, StyleModel] = {}
    try:
        for style in doc.styles:
            if style.type is not None and style.type.name == "PARAGRAPH":
                st = StyleModel()
                f = getattr(style, "font", None)
                if f is not None:
                    st.size = f.size.pt if f.size else None
                    st.bold = f.bold
                    try:
                        rpr = style.element.find(qn("w:rPr"))
                        if rpr is not None:
                            rfonts = rpr.find(qn("w:rFonts"))
                            if rfonts is not None:
                                st.font = rfonts.get(qn("w:ascii"))
                                st.east_asia = rfonts.get(qn("w:eastAsia"))
                    except Exception:
                        pass
                styles[style.name] = st
    except Exception:
        pass
    return styles


def _read_paragraph(p, index: int) -> ParagraphModel:
    pf = p.paragraph_format
    fmt = ParagraphFormatModel()

    if pf.alignment is not None:
        fmt.alignment = _ALIGN_MAP.get(pf.alignment)

    if pf.line_spacing is not None:
        fmt.line_spacing = pf.line_spacing
    if pf.line_spacing_rule is not None:
        fmt.line_spacing_rule = _LINE_RULE_MAP.get(pf.line_spacing_rule)

    if pf.first_line_indent is not None:
        try:
            fmt.first_line_indent = f"{pf.first_line_indent.cm:.2f}cm"
        except Exception:
            fmt.first_line_indent = None
    # 读取 w:firstLineChars，保留“2字符”信息
    first_line_chars = _first_line_chars(p)
    if first_line_chars is not None:
        fmt.first_line_indent = f"{first_line_chars:g}字符"

    if pf.space_before is not None:
        fmt.space_before = pf.space_before.pt
    if pf.space_after is not None:
        fmt.space_after = pf.space_after.pt

    outline_level = _outline_level(p)
    if outline_level is None and p.style is not None:
        outline_level = _style_outline_level(p.style.name)

    runs = [_read_run(r) for r in p.runs]
    return ParagraphModel(
        id=f"p_{index + 1:04d}",
        text=p.text,
        style=p.style.name if p.style is not None else None,
        outline_level=outline_level,
        runs=runs,
        format=fmt,
        has_image=_has_image(p),
        has_formula=_has_formula(p),
    )


def _read_run(r) -> RunModel:
    font = r.font
    east_asia = _run_east_asia(r)
    color = None
    try:
        if font.color is not None and font.color.rgb is not None:
            color = str(font.color.rgb)
    except Exception:
        pass
    return RunModel(
        text=r.text,
        font=font.name,
        east_asia=east_asia,
        size=font.size.pt if font.size else None,
        bold=font.bold,
        italic=font.italic,
        color=color,
    )


def _run_east_asia(r) -> Optional[str]:
    try:
        rpr = r._element.rPr
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                return rfonts.get(qn("w:eastAsia"))
    except Exception:
        pass
    return None


def _outline_level(p) -> Optional[int]:
    try:
        ppr = p._p.pPr
        if ppr is not None and ppr.outlineLvl is not None:
            val = ppr.outlineLvl.get(qn("w:val"))
            if val is not None:
                return int(val)
    except Exception:
        pass
    return None


def _style_outline_level(style_name: Optional[str]) -> Optional[int]:
    if not style_name:
        return None
    for lvl in range(1, 10):
        if style_name in (f"Heading {lvl}", f"标题 {lvl}"):
            return lvl
    return None


def _first_line_chars(p) -> Optional[float]:
    try:
        ppr = p._element.pPr
        if ppr is not None:
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                val = ind.get(qn("w:firstLineChars"))
                if val is not None:
                    return float(val) / 100.0
    except Exception:
        pass
    return None


def _has_image(p) -> bool:
    """检测段落中是否包含嵌入式图片（w:drawing）。"""
    try:
        return p._element.find(qn("w:r")) is not None and p._element.findall(".//" + qn("w:drawing")) != []
    except Exception:
        return False


def _has_formula(p) -> bool:
    """检测段落中是否包含数学公式（MathType OLE 对象或 OMML）。"""
    try:
        # MathType: w:object 内嵌 OLEObject，ProgID 含 Equation
        # o 命名空间不在 python-docx nsmap 中，用完整 URI
        o_ns = "urn:schemas-microsoft-com:office:office"
        for obj in p._element.findall(".//" + qn("w:object")):
            ole = obj.find(f"{{{o_ns}}}OLEObject")
            if ole is not None:
                progid = ole.get("ProgID", "")
                if "Equation" in progid or "MathType" in progid:
                    return True
        # OMML: m:oMath 或 m:oMathPara
        for tag in ("m:oMath", "m:oMathPara"):
            if p._element.findall(".//" + qn(tag)):
                return True
        return False
    except Exception:
        return False
