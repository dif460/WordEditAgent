"""基于 python-docx 的格式化操作（直接修改 docx 对象）。

所有函数均为底层原语，返回 None；由 controller 负责调用并记录变更/支持撤销。
"""
from __future__ import annotations

from typing import Optional

from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

PAGE_SIZES_CM = {
    "A4": (21.0, 29.7),
    "A5": (14.8, 21.0),
    "Letter": (21.59, 27.94),
    "B5": (17.6, 25.0),
}


def set_run_font(
    run,
    font: Optional[str] = None,
    east_asia: Optional[str] = None,
    size: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    f = run.font
    if font is not None:
        f.name = font
    if east_asia is not None:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        f.size = Pt(float(size))
    if bold is not None:
        f.bold = bool(bold)
    if italic is not None:
        f.italic = bool(italic)
    if color is not None:
        from docx.shared import RGBColor

        try:
            f.color.rgb = RGBColor.from_string(color.lstrip("#"))
        except Exception:
            pass


def set_paragraph_alignment(paragraph, alignment: Optional[str]) -> None:
    if alignment and alignment in ALIGN_MAP:
        paragraph.alignment = ALIGN_MAP[alignment]


def set_line_spacing(paragraph, line_spacing: Optional[float]) -> None:
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = float(line_spacing)


def set_space(paragraph, before: Optional[float] = None, after: Optional[float] = None) -> None:
    pf = paragraph.paragraph_format
    if before is not None:
        pf.space_before = Pt(float(before))
    if after is not None:
        pf.space_after = Pt(float(after))


def set_first_line_indent(paragraph, value: Optional[str], pt_size: float = 12.0) -> None:
    """设置首行缩进。value 形如 "2字符" / "0.74cm" / "0字符"。"""
    if not value:
        return
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)

    import re

    v = str(value).strip()
    chars_m = re.search(r"([\d.]+)\s*(?:字符|字)", v)
    cm_m = re.search(r"([\d.]+)\s*cm", v, re.IGNORECASE)
    if chars_m:
        chars = float(chars_m.group(1))
        ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))
        ind.set(qn("w:firstLine"), str(int(round(chars * pt_size * 20))))
    elif cm_m:
        paragraph.paragraph_format.first_line_indent = Cm(float(cm_m.group(1)))
    else:
        # 纯数值按字符处理
        try:
            chars = float(v)
            ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))
            ind.set(qn("w:firstLine"), str(int(round(chars * pt_size * 20))))
        except ValueError:
            return


def apply_style(paragraph, style_name: str) -> None:
    paragraph.style = style_name


def set_outline_level(paragraph, level: Optional[int]) -> None:
    if level is None:
        return
    ppr = paragraph._p.get_or_add_pPr()
    ol = ppr.get_or_add_outlineLvl()
    ol.set(qn("w:val"), str(int(level)))


def modify_style_definition(
    doc,
    style_name: str,
    font: Optional[str] = None,
    east_asia: Optional[str] = None,
    size: Optional[float] = None,
    bold: Optional[bool] = None,
) -> bool:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return False

    if font is not None:
        style.font.name = font
    if size is not None:
        style.font.size = Pt(float(size))
    if bold is not None:
        style.font.bold = bool(bold)
    if east_asia is not None:
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), east_asia)
    return True


def set_section_format(section, page_size: Optional[str] = None, margins: Optional[dict] = None) -> None:
    if page_size:
        size = PAGE_SIZES_CM.get(page_size)
        if size:
            section.page_width = Cm(size[0])
            section.page_height = Cm(size[1])
            section.orientation = WD_ORIENT.PORTRAIT
    if margins:
        if margins.get("top") is not None:
            section.top_margin = Cm(float(margins["top"]))
        if margins.get("bottom") is not None:
            section.bottom_margin = Cm(float(margins["bottom"]))
        if margins.get("left") is not None:
            section.left_margin = Cm(float(margins["left"]))
        if margins.get("right") is not None:
            section.right_margin = Cm(float(margins["right"]))


def set_header_footer(section, header: Optional[str] = None, footer: Optional[str] = None) -> None:
    if header is not None:
        p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        p.text = header
    if footer is not None:
        p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        p.text = footer
