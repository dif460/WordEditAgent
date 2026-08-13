"""文档控制器：加载 docx、执行格式化操作、记录日志并支持撤销。"""
from __future__ import annotations

from io import BytesIO
from typing import Any, Optional

from docx import Document

from engine import formatting as fmt
from engine.reader import build_model
from engine.document_model import DocumentModel


class DocumentController:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        self._history: list[bytes] = []
        self.tool_log: list[dict[str, Any]] = []
        self.changed_ids: list[str] = []

    # ---- 基础能力 ----
    def model(self) -> DocumentModel:
        return build_model(self.document)

    def _snapshot(self) -> None:
        buf = BytesIO()
        self.document.save(buf)
        self._history.append(buf.getvalue())

    def _reload(self, data: bytes) -> None:
        self.document = Document(BytesIO(data))

    def _paragraph_by_id(self, pid: str):
        idx = int(str(pid).split("_")[-1]) - 1
        return self.document.paragraphs[idx]

    def _mark_changed(self, result: dict, pids: list[str]) -> dict:
        for pid in pids:
            if pid not in self.changed_ids:
                self.changed_ids.append(pid)
        result.setdefault("changed", pids)
        result.setdefault("ok", True)
        return result

    def dispatch(self, tool_name: str, args: dict[str, Any]) -> dict[str, Any]:
        method = getattr(self, tool_name, None)
        if method is None:
            return {"ok": False, "changed": [], "message": f"未知工具 {tool_name}"}
        try:
            self._snapshot()
            result = method(**args)
            if not isinstance(result, dict):
                result = {"ok": True, "changed": [], "message": "ok"}
            result.setdefault("ok", True)
            result.setdefault("changed", [])
            self.tool_log.append({"tool": tool_name, "args": args, "result": result})
            return result
        except Exception as e:  # noqa: BLE001
            # 撤销失败的那次快照，保持文档与历史一致
            if self._history:
                self._reload(self._history.pop())
            return {"ok": False, "changed": [], "message": f"{tool_name} 执行失败: {e}"}

    # ---- 读取类工具 ----
    def read_document(self) -> dict[str, Any]:
        return {"ok": True, "model": self.model().model_dump(), "message": "已读取文档"}

    def get_document_overview(self) -> dict[str, Any]:
        return {"ok": True, "overview": self.model().overview(), "message": "已生成结构摘要"}

    def get_paragraph_detail(self, paragraph_id: str) -> dict[str, Any]:
        p = self._paragraph_by_id(paragraph_id)
        from engine.reader import _read_paragraph  # noqa

        idx = int(paragraph_id.split("_")[-1]) - 1
        detail = _read_paragraph(p, idx).model_dump()
        return {"ok": True, "paragraph": detail, "message": f"已查看 {paragraph_id}"}

    # ---- 格式化类工具 ----
    def set_heading_style(
        self,
        paragraph_id: str,
        level: int,
        font: Optional[str] = None,
        east_asia: Optional[str] = None,
        size: Optional[float] = None,
        bold: Optional[bool] = None,
        alignment: Optional[str] = None,
    ) -> dict[str, Any]:
        p = self._paragraph_by_id(paragraph_id)
        style_name = f"Heading {level}"
        if style_name in [s.name for s in self.document.styles]:
            p.style = self.document.styles[style_name]
        fmt.set_outline_level(p, level)

        ea = east_asia or font
        for r in p.runs:
            fmt.set_run_font(r, font=font, east_asia=ea, size=size, bold=bold)
        if alignment:
            fmt.set_paragraph_alignment(p, alignment)
        return self._mark_changed(
            {"message": f"{paragraph_id} 已设为 {style_name}：{font or ''} {size or ''}pt"},
            [paragraph_id],
        )

    def set_paragraph_format(
        self,
        paragraph_id: str,
        alignment: Optional[str] = None,
        line_spacing: Optional[float] = None,
        first_line_indent: Optional[str] = None,
        space_before: Optional[float] = None,
        space_after: Optional[float] = None,
    ) -> dict[str, Any]:
        p = self._paragraph_by_id(paragraph_id)
        if alignment:
            fmt.set_paragraph_alignment(p, alignment)
        if line_spacing is not None:
            fmt.set_line_spacing(p, line_spacing)
        if first_line_indent is not None:
            size = _para_size_pt(p)
            fmt.set_first_line_indent(p, first_line_indent, pt_size=size)
        if space_before is not None or space_after is not None:
            fmt.set_space(p, before=space_before, after=space_after)
        return self._mark_changed(
            {"message": f"{paragraph_id} 段落格式已更新"}, [paragraph_id]
        )

    def set_run_font(
        self,
        paragraph_id: str,
        font: Optional[str] = None,
        east_asia: Optional[str] = None,
        size: Optional[float] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        color: Optional[str] = None,
        run_index: Optional[int] = None,
    ) -> dict[str, Any]:
        p = self._paragraph_by_id(paragraph_id)
        runs = p.runs
        if run_index is not None:
            runs = [runs[run_index]]
        for r in runs:
            fmt.set_run_font(r, font=font, east_asia=east_asia, size=size, bold=bold, italic=italic, color=color)
        return self._mark_changed(
            {"message": f"{paragraph_id} 字体已更新为 {east_asia or font or ''} {size or ''}pt"},
            [paragraph_id],
        )

    def apply_style(self, paragraph_id: str, style_name: str) -> dict[str, Any]:
        p = self._paragraph_by_id(paragraph_id)
        fmt.apply_style(p, style_name)
        return self._mark_changed({"message": f"{paragraph_id} 已应用样式 {style_name}"}, [paragraph_id])

    def modify_style_definition(
        self,
        style_name: str,
        font: Optional[str] = None,
        east_asia: Optional[str] = None,
        size: Optional[float] = None,
        bold: Optional[bool] = None,
    ) -> dict[str, Any]:
        ok = fmt.modify_style_definition(
            self.document, style_name, font=font, east_asia=east_asia, size=size, bold=bold
        )
        return {
            "ok": ok,
            "changed": [] if not ok else [style_name],
            "message": f"样式 {style_name} 已修改" if ok else f"样式 {style_name} 不存在",
        }

    def set_section_format(
        self,
        page_size: Optional[str] = None,
        margins: Optional[dict] = None,
        section_index: int = 0,
    ) -> dict[str, Any]:
        section = self.document.sections[section_index]
        fmt.set_section_format(section, page_size=page_size, margins=margins)
        return {"ok": True, "changed": [], "message": "页面设置已更新"}

    def set_header_footer(self, header: Optional[str] = None, footer: Optional[str] = None) -> dict[str, Any]:
        for section in self.document.sections:
            fmt.set_header_footer(section, header=header, footer=footer)
        return {"ok": True, "changed": [], "message": "页眉页脚已更新"}

    def set_table_font(
        self,
        table_id: str,
        font: Optional[str] = None,
        east_asia: Optional[str] = None,
        size: Optional[float] = None,
    ) -> dict[str, Any]:
        idx = int(str(table_id).split("_")[-1]) - 1
        table = self.document.tables[idx]
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        fmt.set_run_font(r, font=font, east_asia=east_asia, size=size)
        return self._mark_changed({"message": f"表格 {table_id} 字体已更新"}, [table_id])

    def update_toc(self) -> dict[str, Any]:
        """用 Word COM 刷新目录字段（Windows 专用，尽力而为）。"""
        import os
        import tempfile

        from engine import word_com

        tmp = os.path.join(tempfile.gettempdir(), f"wordedit_toc_{os.getpid()}.docx")
        try:
            self.document.save(tmp)
            res = word_com.update_toc(tmp)
            if res.get("ok"):
                with open(tmp, "rb") as f:
                    self._reload(f.read())
            return {"ok": bool(res.get("ok")), "changed": [], "message": res.get("message", "")}
        finally:
            if os.path.exists(tmp):
                os.remove(tmp)

    # ---- 交付类 ----
    def save_document(self, output_path: str) -> dict[str, Any]:
        self.document.save(output_path)
        return {"ok": True, "changed": [], "message": f"已保存 {output_path}", "output_path": output_path}

    def undo_last(self) -> dict[str, Any]:
        if not self._history:
            return {"ok": False, "changed": [], "message": "没有可撤销的步骤"}
        self._reload(self._history.pop())
        if self.tool_log:
            self.tool_log.pop()
        return {"ok": True, "changed": [], "message": "已回滚上一步"}


def _para_size_pt(p) -> float:
    for r in p.runs:
        if r.font.size is not None:
            return r.font.size.pt
    try:
        st = p.style
        if st is not None and st.font.size is not None:
            return st.font.size.pt
    except Exception:
        pass
    return 12.0
