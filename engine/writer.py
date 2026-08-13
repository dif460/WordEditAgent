"""DocumentModel -> docx 写入（用于按模型重建文档）。"""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from engine.document_model import DocumentModel
from engine import formatting as fmt


def write_document(model: DocumentModel, output_path: str) -> None:
    doc = Document()

    if model.sections:
        sec = model.sections[0]
        fmt.set_section_format(doc.sections[0], page_size=sec.page_size, margins=sec.margins)
        if sec.header:
            doc.sections[0].header.paragraphs[0].text = sec.header
        if sec.footer:
            doc.sections[0].footer.paragraphs[0].text = sec.footer

    for pm in model.paragraphs:
        p = doc.add_paragraph()
        if pm.style:
            try:
                p.style = pm.style
            except Exception:
                pass
        fmt.set_outline_level(p, pm.outline_level)
        pf = pm.format
        if pf.alignment:
            fmt.set_paragraph_alignment(p, pf.alignment)
        if pf.line_spacing:
            fmt.set_line_spacing(p, pf.line_spacing)
        if pf.first_line_indent:
            fmt.set_first_line_indent(p, pf.first_line_indent)
        if pf.space_before is not None or pf.space_after is not None:
            fmt.set_space(p, before=pf.space_before, after=pf.space_after)

        if pm.runs:
            for rm in pm.runs:
                run = p.add_run(rm.text)
                fmt.set_run_font(
                    run,
                    font=rm.font,
                    east_asia=rm.east_asia,
                    size=rm.size,
                    bold=rm.bold,
                    italic=rm.italic,
                    color=rm.color,
                )
        else:
            p.add_run(pm.text)

    doc.save(output_path)
