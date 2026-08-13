"""生成示例中文 docx，用于本地测试。"""
from __future__ import annotations

from pathlib import Path

from docx import Document


def generate_sample(path: str) -> None:
    doc = Document()

    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph(
        "随着信息技术的不断发展，文档的自动化处理需求日益增长。"
        "本文针对 Word 文档的格式规范化问题展开研究，提出一套智能化的格式编辑方案。"
    )
    doc.add_heading("1.1 研究背景", level=2)
    doc.add_paragraph(
        "在日常办公中，大量的文档需要遵循统一的排版规范。传统的人工排版方式效率低下，"
        "且容易出错。因此，利用人工智能技术实现文档的自动格式化具有重要的现实意义。"
    )
    doc.add_heading("1.2 研究内容", level=2)
    doc.add_paragraph(
        "本研究主要包括需求解析、文档建模、格式规则生成、格式执行与校验等核心环节。"
    )
    doc.add_heading("第二章 系统设计", level=1)
    doc.add_paragraph(
        "系统采用分层架构，前端负责文档导入与预览，后端负责格式化的编排与执行。"
    )

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "模块"
    table.cell(0, 1).text = "功能"
    table.cell(1, 0).text = "引擎层"
    table.cell(1, 1).text = "文档读写与格式化"

    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"已生成示例文档: {path}")


if __name__ == "__main__":
    from app.config import settings

    generate_sample(str(settings.upload_path / "sample.docx"))
